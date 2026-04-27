import os, sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers), style='Light Grid Accent 1')
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, rd in enumerate(rows):
        for ci, cd in enumerate(rd):
            cell = t.rows[ri+1].cells[ci]
            cell.text = cd
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    if ci == 1 and cd in ('YAPILDI', 'TAMAMLANDI'):
                        r.bold = True
                        r.font.color.rgb = RGBColor(0,128,0)
                    elif ci == 1 and cd in ('YAPILMADI', 'BEKLIYOR'):
                        r.bold = True
                        r.font.color.rgb = RGBColor(255,0,0)
    return t

# ===== KAPAK =====
doc.add_paragraph('')
doc.add_paragraph('')
title = doc.add_heading('Coachify Mentor Payment Panel', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_heading('Yapilan Degisiklikler Raporu', level=1)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')
doc.add_paragraph('Tarih: 21 Nisan 2026').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Revize Kaynak: Coachify 2. Revize-2.pdf').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')
doc.add_paragraph('')

# ===== GENEL OZET =====
doc.add_heading('1. Genel Ozet', level=1)
doc.add_paragraph(
    'Bu belge, Coachify Mentor Payment Panel uzerinde yapilan tum revize degisikliklerini '
    'detayli sekilde aciklamaktadir. Degisiklikler, musteri tarafindan gonderilen '
    '"2. Revize-2" belgesindeki 9 madde uzerinden degerlendirilmis ve uygulanmistir.'
)

doc.add_paragraph('')
add_table(doc,
    ['No', 'Revize Maddesi', 'Durum'],
    [
        ['1', 'Donem tarihi 1 ve 15 olmali (14 ve 15 gorunuyordu)', 'YAPILDI'],
        ['2', 'Yeni orenci formunda okul kaldirilmali, * isaretleri duzeltilmeli', 'YAPILDI'],
        ['3', 'Yeni orenci odeme durumu "odendi" baslamali', 'YAPILDI (zaten doru)'],
        ['4', 'Orenci bitis tarihi aylik dongu ile hesaplanmali (haftalik degil)', 'YAPILDI'],
        ['5', 'Mentor maksimum 3 odeme gormeli, paket suresi gorunmemeli', 'YAPILDI'],
        ['6', 'Her orencinin bitis tarihi zorunlu olmali', 'YAPILDI (zaten doru)'],
        ['7', 'Iade hesabi kontrolu (tamamlanan hafta kadar odeme)', 'YAPILDI (zaten doru)'],
        ['8', 'Webhook baglantisi ve Zapier entegrasyonu', 'YAPILDI (zaten mevcut)'],
        ['9', 'Odeme bilgilendirme mesajlari', 'YAPILMADI'],
    ]
)

doc.add_paragraph('')
doc.add_paragraph(
    'Not: 3, 6, 7 ve 8. maddeler mevcut sistemde zaten doru sekilde calisiyordu. '
    '1, 2, 4 ve 5 numarali maddeler icin kod degisiklii yapildi. '
    '9. madde (bildirim mesajlari) Zapier + CRM entegrasyonu ile disaridan cozulecektir.'
)

# ===== MADDE 1 =====
doc.add_heading('2. Madde 1 \u2014 Donem Tarihi Gosterimi', level=1)

doc.add_heading('2.1 Sorun', level=2)
doc.add_paragraph(
    'Bazi yerlerde donem tarihi 14 ve 15 olarak gorunuyordu. '
    'Mantigin 1 ve 15 baz almasi gerekiyordu.'
)

doc.add_heading('2.2 Yapilan Degisiklik', level=2)
doc.add_paragraph(
    'Donem tarihi hesaplama sistemi temelden degistirildi. '
    'Sabit 1/15 dongusu yerine, ogrencinin SAG (Satin Alma Gunu) bazli aylik dongu sistemine gecildi. '
    'Bu sayede donem tarihleri ogrenciye ozel hesaplanirken, odeme gunleri 1 ve 15 olarak sabit kaldi.'
)

doc.add_heading('2.3 Yeni Sistem Mantii', level=2)
doc.add_paragraph('Iki ayri kavram net olarak ayrildi:')

add_table(doc,
    ['Kavram', 'Aciklama', 'Ornek'],
    [
        ['Dongu Tarihi (cycleDate)', 'SAG bazli aylik. Ogrencinin baslangic gunu + N ay', 'SAG 4 Nisan \u2192 4 Mayis, 4 Haziran, 4 Temmuz'],
        ['Odeme Tarihi (paymentDate)', 'Dongu tarihinden hesaplanan odeme gunu (1 veya 15)', '4 Mayis \u2192 15 Mayis | 19 Mayis \u2192 1 Haziran'],
    ]
)

doc.add_paragraph('')
doc.add_paragraph('Odeme tarihi hesaplama kurali:')
doc.add_paragraph('Dongu tarihinin gunu 15 veya kucukse \u2192 Ayni ayin 15\'inde odeme', style='List Bullet')
doc.add_paragraph('Dongu tarihinin gunu 15\'ten buyukse \u2192 Sonraki ayin 1\'inde odeme', style='List Bullet')

doc.add_heading('2.4 Ornek Senaryolar', level=2)
add_table(doc,
    ['SAG', 'Dongu Tarihleri', 'Odeme Tarihleri'],
    [
        ['2 Ocak', '2 Subat, 2 Mart, 2 Nisan, 2 Mayis', '15 Subat, 15 Mart, 15 Nisan, 15 Mayis'],
        ['10 Nisan', '10 Mayis, 10 Haziran, 10 Temmuz', '15 Mayis, 15 Haziran, 15 Temmuz'],
        ['19 Nisan', '19 Mayis, 19 Haziran, 19 Temmuz', '1 Haziran, 1 Temmuz, 1 Agustos'],
        ['1 Subat', '1 Mart, 1 Nisan, 1 Mayis', '15 Mart, 15 Nisan, 15 Mayis'],
    ]
)

doc.add_heading('2.5 Degisen Dosyalar', level=2)
add_table(doc,
    ['Dosya', 'Degisiklik'],
    [
        ['src/lib/mentor-earnings.ts', 'getAllCycleDates fonksiyonu yeniden yazildi (SAG bazli aylik dongu)'],
        ['src/lib/mentor-earnings.ts', 'getNextPaymentDate fonksiyonu yeniden yazildi (SAG bazli)'],
        ['src/lib/mentor-earnings.ts', 'Yeni getPaymentDateForCycle fonksiyonu eklendi (1/15 hesaplama)'],
        ['src/app/api/mentor/earnings/route.ts', 'Odeme tarihi bazli gruplama ve 3 dongu siniri'],
        ['src/app/mentor/page.tsx', 'paymentDate ile gosterim ve gruplama'],
    ]
)

doc.add_heading('2.6 Eski vs Yeni Kod Karsilastirmasi', level=2)
doc.add_paragraph('Eski getAllCycleDates (50+ satir, sabit 1/15 mantii):')
p = doc.add_paragraph()
p.style = doc.styles['Normal']
run = p.add_run(
    'UBG 1-15 ise: 15, 1, 15, 1, 15, 1...\n'
    'UBG 16+ ise: 1, 15, 1, 15, 1, 15...'
)
run.font.size = Pt(9)
run.font.name = 'Consolas'

doc.add_paragraph('Yeni getAllCycleDates (10 satir, SAG bazli aylik):')
p = doc.add_paragraph()
run = p.add_run(
    'SAG + 1 ay, SAG + 2 ay, SAG + 3 ay...\n'
    'Ornek: 4 Nisan \u2192 4 Mayis, 4 Haziran, 4 Temmuz...'
)
run.font.size = Pt(9)
run.font.name = 'Consolas'

# ===== MADDE 2 =====
doc.add_heading('3. Madde 2 \u2014 Yeni Orenci Ekleme Formu', level=1)

doc.add_heading('3.1 Sorun', level=2)
doc.add_paragraph('Okul alani formda yer aliyordu ama kayitta okul bilgisi toplanmiyordu. '
    'Zorunlu ve opsiyonel alanlar net ayrilmamisti. * isaretleri eksikti.')

doc.add_heading('3.2 Yapilan Degisiklik', level=2)
add_table(doc,
    ['Alt Istek', 'Durum', 'Detay'],
    [
        ['Email zorunlu olmamali', 'YAPILDI', 'Formda "required" yok, API null kabul ediyor'],
        ['Okul alani kaldirilmali', 'YAPILDI', 'Formdan zaten yoktu. API ve detay sayfasindan kaldirildi'],
        ['Zorunlu alanlar * ile isaretli', 'YAPILDI', '6 zorunlu alanin hepsinde * mevcut'],
        ['Opsiyonel alanlar etiketli', 'YAPILDI', '"(Opsiyonel)" etiketi eklendi'],
    ]
)

doc.add_heading('3.3 Form Alani Durumu', level=2)
add_table(doc,
    ['Alan', 'Tip', 'Isaret'],
    [
        ['Ad Soyad', 'Zorunlu', '*'],
        ['Telefon', 'Zorunlu', '*'],
        ['Mentor', 'Zorunlu', '*'],
        ['Baslangic Tarihi', 'Zorunlu', '*'],
        ['Paket Turu', 'Zorunlu', '*'],
        ['Bitis Tarihi', 'Zorunlu', '*'],
        ['E-posta', 'Opsiyonel', '(Opsiyonel)'],
        ['Sinif', 'Opsiyonel', '(Opsiyonel)'],
        ['Veli Bilgileri', 'Opsiyonel', '(Opsiyonel)'],
        ['Iletisim Tercihi', 'Opsiyonel', '(Opsiyonel)'],
        ['Indirim Kodu', 'Opsiyonel', '(Opsiyonel)'],
        ['Ozel Aciklama', 'Opsiyonel', '(Opsiyonel)'],
    ]
)

doc.add_heading('3.4 Degisen Dosyalar', level=2)
add_table(doc,
    ['Dosya', 'Degisiklik'],
    [
        ['src/app/admin/students/[id]/page.tsx', 'Okul (school) alani gosterimi kaldirildi'],
        ['src/app/api/admin/students/route.ts', 'school destructuring kaldirildi, school: "" sabiti'],
        ['src/app/api/admin/students/[id]/route.ts', 'school destructuring ve updateData kaldirildi'],
        ['src/app/admin/students/page.tsx', '(Opsiyonel) etiketleri eklendi'],
    ]
)

# ===== MADDE 4 =====
doc.add_heading('4. Madde 4 \u2014 Aylik Dongu ile Bitis Tarihi Hesabi', level=1)

doc.add_heading('4.1 Sorun', level=2)
doc.add_paragraph(
    'Orenci bitis tarihleri haftalik mantikla (1 ve 15\'e sabit) ilerliyordu. '
    'Brief kural 4.5\'e gore ogrencinin baslangic gununu baz alan aylik dongu ile ilerlemeli.'
)

doc.add_heading('4.2 Yapilan Degisiklik', level=2)
doc.add_paragraph(
    'Bu degisiklik Madde 1 ile ayni kod tabanini kapsar. '
    'Temel degisiklik mentor-earnings.ts dosyasindaki dongu hesaplama fonksiyonlaridir. '
    'Ayrica bu degisiklik icin veri gecisi yapildi:'
)

doc.add_paragraph('')
doc.add_paragraph('Veri Gecisi Adimlari:', style='List Bullet')
doc.add_paragraph('Mevcut periodic_calc pending earnings iptal edildi (SQL ile)', style='List Bullet 2')
doc.add_paragraph('Admin panelde "Hakedis Hesapla" tetiklenerek yeni SAG bazli earnings olusturuldu', style='List Bullet 2')
doc.add_paragraph('Paid ve cancelled kayitlara dokunulmadi (kural 4.8)', style='List Bullet 2')

doc.add_heading('4.3 Dogrulama', level=2)
doc.add_paragraph(
    'Yeni sistemde veritabanina yazilan MentorEarning kayitlarinin cycleDate degerleri '
    'ogrenci baslangic gunu bazli olarak dogrulanmistir:'
)

add_table(doc,
    ['Orenci', 'SAG', 'Yeni Cycle Date\'ler', 'Odeme Tarihleri'],
    [
        ['Idil Demirci', '2 Ocak', '2 Sub, 2 Mar, 2 Nis', '15 Sub, 15 Mar, 15 Nis'],
        ['Betul Kaya', '3 Ocak', '3 Subat', '15 Subat'],
        ['Gizem Sever', '2 Eylul', '2 Eki, 2 Kas, 2 Ara, 2 Oca, 2 Sub, 2 Mar, 2 Nis', '15\'ler'],
        ['Kardelen Cakmak', '16 Nisan', '16 Mayis (gelecek)', '1 Haziran'],
        ['Ilknur Turan', '19 Nisan', '19 Mayis (gelecek)', '1 Haziran'],
    ]
)

# ===== MADDE 5 =====
doc.add_heading('5. Madde 5 \u2014 Mentor Odeme Gorunurlueu', level=1)

doc.add_heading('5.1 Sorun', level=2)
doc.add_paragraph(
    'Mentor, ogrenciden ne kadar sure odeme alacaini gorebiliyordu. '
    'Bu sayede dolayli olarak ogrencinin kac aylik paket aldiini anlayabiliyordu. '
    'Mentor maksimum sadece onundeki 3 odemeyi gormeliydi.'
)

doc.add_heading('5.2 Yapilan Degisiklik', level=2)
doc.add_paragraph(
    'Mentor API route\'unda (api/mentor/earnings) 3 dongu siniri uygulanmistir. '
    'Sinirlama ODEME TARIHI bazli yapilir (1/15), dongu tarihi bazli degil. '
    'Bu fark onemlidir: farkli SAG\'lara sahip ogrenciler ayni odeme tarihine sahip olabilir.'
)

doc.add_paragraph('')
doc.add_paragraph('Ornek: Mentorun 3 ogrencisi varsa:', style='List Bullet')
doc.add_paragraph('Ogrenci A (SAG 4) \u2192 cycle 4 Mayis \u2192 odeme 15 Mayis', style='List Bullet 2')
doc.add_paragraph('Ogrenci B (SAG 10) \u2192 cycle 10 Mayis \u2192 odeme 15 Mayis', style='List Bullet 2')
doc.add_paragraph('Ogrenci C (SAG 20) \u2192 cycle 20 Mayis \u2192 odeme 1 Haziran', style='List Bullet 2')
doc.add_paragraph('Mentor 2 odeme tarihi gorur: 15 Mayis ve 1 Haziran', style='List Bullet 2')

add_table(doc,
    ['Bilgi', 'Mentor Gorur mu?'],
    [
        ['Onundeki 3 odeme tarihi', 'EVET'],
        ['3\'u asan gelecek odemeler', 'HAYIR'],
        ['Ogrencinin toplam paket suresi', 'HAYIR'],
        ['Ogrencinin kac aylik paket aldii', 'HAYIR'],
        ['Geamis tum hakedisleri', 'EVET'],
    ]
)

doc.add_heading('5.3 Teknik Detay', level=2)
doc.add_paragraph('Mentor API sinirlama mantii:')
doc.add_paragraph('pending earnings\'lerin cycleDate\'lerinden paymentDate hesaplanir', style='List Bullet 2')
doc.add_paragraph('Benzersiz paymentDate\'ler bulunur, kronolojik siralanir', style='List Bullet 2')
doc.add_paragraph('Ilk 3 odeme tarihine denk gelen earnings\'ler gosterilir', style='List Bullet 2')
doc.add_paragraph('totalPending sadece gorunen earnings\'lerin toplami olarak hesaplanir', style='List Bullet 2')

# ===== MADDE 9 =====
doc.add_heading('6. Madde 9 \u2014 Odeme Bilgilendirme Mesajlari (Bekleyen)', level=1)

doc.add_heading('6.1 Durum', level=2)
doc.add_paragraph(
    'Bu madde panel tarafinda yapilmasi gereken bir kod degisikliii icermemektedir. '
    'Mevcut sistemde /api/webhooks/zapier/query endpoint\'i "expiring_soon" sorgusu ile '
    'bitisina yaklasan orencileri listelemeye hazirdir.'
)

doc.add_heading('6.2 Gerekli Zapier Akisi', level=2)
doc.add_paragraph('Zapier\'de kurulmasi gereken akis:')
doc.add_paragraph('')
doc.add_paragraph('1. Trigger: Her gun saat 17:00', style='List Number')
doc.add_paragraph('2. Action: HTTP Request \u2192 POST /api/webhooks/zapier/query', style='List Number')
p = doc.add_paragraph()
p.style = 'List Bullet 2'
run = p.add_run('{ "query": "expiring_soon", "filters": { "daysAhead": 3 } }')
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph('3. Action: Her orenci icin Kommo CRM\'de kisi bul', style='List Number')
doc.add_paragraph('4. Action: CRM status\'u "hatirlatma" olarak guncelle', style='List Number')
doc.add_paragraph('5. CRM otomatik mesaj gonderir', style='List Number')

doc.add_paragraph('')
doc.add_paragraph('Query endpoint\'in dondurdueu veriler:', style='List Bullet')
doc.add_paragraph('name, phone, email, endDate, daysUntilExpiry', style='List Bullet 2')
doc.add_paragraph('parentName, parentPhone, contactPreference', style='List Bullet 2')
doc.add_paragraph('currentMentor, membershipType', style='List Bullet 2')

# ===== TOPLAM DEGISIKLIK OZETI =====
doc.add_heading('7. Toplam Degisiklik Ozeti', level=1)

doc.add_heading('7.1 Degisen Dosyalar', level=2)
add_table(doc,
    ['Dosya Yolu', 'Degisiklik Turu', 'Aciklama'],
    [
        ['src/lib/mentor-earnings.ts', 'Yeniden Yazildi', 'getAllCycleDates, getNextPaymentDate yeniden yazildi. getPaymentDateForCycle eklendi.'],
        ['src/app/api/mentor/earnings/route.ts', 'Degistirildi', 'Odeme tarihi bazli gruplama, 3 dongu siniri, paymentDate alani'],
        ['src/app/mentor/page.tsx', 'Degistirildi', 'paymentDate ile gruplama ve gosterim. Earning interface guncellendi.'],
        ['src/app/admin/students/page.tsx', 'Degistirildi', '(Opsiyonel) etiketleri eklendi'],
        ['src/app/admin/students/[id]/page.tsx', 'Degistirildi', 'Okul alani gosterimi kaldirildi'],
        ['src/app/api/admin/students/route.ts', 'Degistirildi', 'school alani temizlendi'],
        ['src/app/api/admin/students/[id]/route.ts', 'Degistirildi', 'school alani temizlendi'],
    ]
)

doc.add_heading('7.2 Dokunulmayan (Zaten Doru Calisan) Kisimlar', level=2)
add_table(doc,
    ['Kisim', 'Aciklama'],
    [
        ['Haftalik tutar hesabi', '375 TL/hafta hesabi doru, degisiklik yapilmadi (kural 4.2)'],
        ['Yeni orenci odeme durumu', 'paymentStatus default "paid" (kural 3)'],
        ['Bitis tarihi zorunlulueu', 'endDate DateTime (nullable degil) (kural 6)'],
        ['Iade hesabi', 'Tamamlanan hafta kadar odeme (kural 4.7)'],
        ['Duplicate odeme koruma', 'upsert + unique constraint (kural 4.8)'],
        ['Webhook endpoint\'leri', '4 Zapier endpoint mevcut ve calisir (kural 8)'],
        ['Rate-lock mekanizmasi', 'Her earning kaydinda weeklyRate sabitleniyor'],
    ]
)

doc.add_heading('7.3 Veri Gecisi', level=2)
doc.add_paragraph(
    'Eski sabit 1/15 dongu tarihlerine sahip periodic_calc pending earnings iptal edildi. '
    'Yeni SAG bazli aylik dongu tarihleri ile yeniden hesaplandi. '
    'Paid ve cancelled kayitlara dokunulmadi. '
    'Prisma schema degisikliii yapilmadi.'
)

# Kaydet
path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Coachify_Degisiklik_Raporu.docx')
doc.save(path)
print('Dosya olusturuldu: ' + path)
