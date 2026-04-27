import os, sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

def add_table(doc, headers, rows):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers), style='Light Grid Accent 1')
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, rd in enumerate(rows):
        for ci, cd in enumerate(rd):
            t.rows[ri+1].cells[ci].text = cd
            for p in t.rows[ri+1].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return t

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_bold_paragraph(doc, bold_text, normal_text=''):
    p = doc.add_paragraph()
    run = p.add_run(bold_text)
    run.bold = True
    if normal_text:
        p.add_run(normal_text)
    return p

# ===== KAPAK =====
doc.add_paragraph('')
doc.add_paragraph('')
title = doc.add_heading('Coachify Mentor Odeme Mantigi', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_heading('Aylik / Haftalik Kayit ve Odeme Sistemi Detayli Aciklamasi', level=2)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')
doc.add_paragraph('Tarih: 22 Nisan 2026').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Kaynak: Proje Briefi + 2. Revize + Revize Raporu')
doc.add_paragraph('')
doc.add_paragraph('')

# ===== 1. GENEL BAKIS =====
add_heading(doc, '1. Genel Bakis', level=1)

doc.add_paragraph(
    'Coachify\'da mentor odeme sistemi iki temel kavram uzerine kuruludur: '
    'AYLIK odeme ve HAFTALIK odeme. Bu iki mod farkli senaryolarda kullanilir '
    've projenin en kritik is mantigini olusturur.'
)

doc.add_paragraph(
    'Temel kural (Brief 4.1): "Her mentor, ogrenci basina aylik 1500 TL kazanir."'
)

add_table(doc,
    ['Hesaplama Modu', 'Birim Tutar', 'Kullanildigi Yer', 'Brief Kurali'],
    [
        ['AYLIK', '1500 TL / ay', 'Periyodik hesaplama (normal durum)', 'Kural 4.1, 4.5'],
        ['HAFTALIK', '375 TL / hafta', 'Sonlandirma (mentor degisikligi, iade, birakma)', 'Kural 4.2, 4.6, 4.7'],
    ]
)

doc.add_paragraph('')
doc.add_paragraph('Iki mod arasindaki iliski:', style='List Bullet')
doc.add_paragraph('Aylik 1500 TL = 4 hafta x 375 TL (haftalik bolunmus hali)', style='List Bullet 2')
doc.add_paragraph('Aylik hesaplama tam ay icin kullanilir, haftalik ise kismi sureler icin', style='List Bullet 2')
doc.add_paragraph('Bu iki mod ayni anda, ayni ogrenci icin farkli zamanlarda kullanilabilir', style='List Bullet 2')

# ===== 2. TEMEL KAVRAMLAR =====
add_heading(doc, '2. Temel Kavramlar', level=1)

add_heading(doc, '2.1 SAG (Satin Alma Gunu)', level=2)
doc.add_paragraph(
    'Ogrencinin sisteme kayit oldugu tarih. Odeme dongusunun baslangic noktasidir. '
    'Bu tarih degismez (Brief 3.3).'
)
doc.add_paragraph('Ornek: Ogrenci 10 Nisan\'da kayit oldu → SAG = 10 Nisan', style='List Bullet')

add_heading(doc, '2.2 Dongu Tarihi (Cycle Date)', level=2)
doc.add_paragraph(
    'Her tamamlanan ayin isaretlendigi tarih. SAG bazinda aylik olarak hesaplanir (Kural 4.5).'
)
doc.add_paragraph('SAG 10 Nisan → Dongu tarihleri: 10 Mayis, 10 Haziran, 10 Temmuz, 10 Agustos...', style='List Bullet')
doc.add_paragraph(
    'Her dongu tarihi geldiginde, bir onceki dongu tarihinden (veya SAG\'dan) bu dongu tarihine '
    'kadar gecen 1 aylik sure icin mentor 1500 TL kazanir.'
)

add_heading(doc, '2.3 Odeme Tarihi (Payment Date)', level=2)
doc.add_paragraph(
    'Mentorun parayi hesabina alacagi gun. Dongu tarihinden hesaplanir, her zaman 1 veya 15\'e denk gelir (Kural 4.3).'
)

add_table(doc,
    ['Dongu Tarihi Gunu', 'Odeme Tarihi', 'Ornek'],
    [
        ['1 - 15 arasi', 'Ayni ayin 15\'i', 'Dongu 10 Mayis → Odeme 15 Mayis'],
        ['16 ve sonrasi', 'Sonraki ayin 1\'i', 'Dongu 19 Mayis → Odeme 1 Haziran'],
    ]
)

doc.add_paragraph('')
doc.add_paragraph('Onemli: Dongu tarihi ogrencinin ilerleme gostergesidir. Odeme tarihi ise para transfer gunudur. Bu ikisi farkli tarihler olabilir.')

# ===== 3. AYLIK HESAPLAMA DETAYI =====
add_heading(doc, '3. Mod 1: Aylik Hesaplama (Periyodik)', level=1)

add_heading(doc, '3.1 Ne Zaman Kullanilir?', level=2)
doc.add_paragraph('Ogrenci aktif olarak devam ettigi surece, her ayin sonunda otomatik olarak calisir. Bu "periyodik hesaplama" (periodic_calc) olarak adlandirilir.')

add_heading(doc, '3.2 Mantik', level=2)
doc.add_paragraph(
    'Her dongu tarihi geldiginde, o donem icin sabit 1500 TL yazilir. '
    'Kac gun oldugu onemli degil — 28 gun de olsa 31 gun de olsa, tamamlanan her ay = 1500 TL.'
)

add_heading(doc, '3.3 Adim Adim Ornek', level=2)

add_bold_paragraph(doc, 'Ogrenci: Ahmet, SAG: 10 Nisan, Paket: 3 Aylik')

doc.add_paragraph('')
doc.add_paragraph('Dongu tarihleri ve odemeler:')

add_table(doc,
    ['Donem', 'Baslangic', 'Bitis (Dongu)', 'Gecen Gun', 'Tamamlanan Ay', 'Mentor Kazanci', 'Odeme Tarihi'],
    [
        ['1. ay', '10 Nisan', '10 Mayis', '30 gun', '1 ay', '1500 TL', '15 Mayis'],
        ['2. ay', '10 Mayis', '10 Haziran', '31 gun', '1 ay', '1500 TL', '15 Haziran'],
        ['3. ay', '10 Haziran', '10 Temmuz', '30 gun', '1 ay', '1500 TL', '15 Temmuz'],
    ]
)

doc.add_paragraph('')
add_bold_paragraph(doc, 'Toplam: 3 ay x 1500 TL = 4500 TL')
doc.add_paragraph('')

add_heading(doc, '3.4 Zamanlama Gorseli', level=2)

p = doc.add_paragraph()
run = p.add_run(
    'Nisan           Mayis           Haziran         Temmuz\n'
    '|----10----|----10----|----10----|----10---->  (dongu tarihleri)\n'
    '|  1500 TL  |  1500 TL  |  1500 TL  |\n'
    '\n'
    'Odeme gunleri:\n'
    '              15 Mayis      15 Haziran     15 Temmuz\n'
    '              (1.odeme)     (2.odeme)      (3.odeme)'
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

add_heading(doc, '3.5 Farkli SAG Ornekleri', level=2)

add_table(doc,
    ['SAG', '1. Dongu', '1. Odeme', '2. Dongu', '2. Odeme', 'Neden?'],
    [
        ['2 Ocak', '2 Subat', '15 Subat', '2 Mart', '15 Mart', '2 <= 15 → 15\'inde'],
        ['10 Nisan', '10 Mayis', '15 Mayis', '10 Haziran', '15 Haziran', '10 <= 15 → 15\'inde'],
        ['15 Mart', '15 Nisan', '15 Nisan', '15 Mayis', '15 Mayis', '15 <= 15 → 15\'inde'],
        ['16 Subat', '16 Mart', '1 Nisan', '16 Nisan', '1 Mayis', '16 > 15 → sonraki ayin 1\'i'],
        ['19 Nisan', '19 Mayis', '1 Haziran', '19 Haziran', '1 Temmuz', '19 > 15 → sonraki ayin 1\'i'],
        ['31 Ocak', '28 Subat', '15 Subat', '31 Mart', '15 Mart', 'Ay sonu tasmasi: 31→28'],
    ]
)

add_heading(doc, '3.6 Artan (Incremental) Hesaplama', level=2)
doc.add_paragraph(
    'Sistem her dongu tarihi icin kumulatif toplam hesaplar ve onceki donemlerden kalani cikarir. '
    'Bu sayede duplicate odeme onlenir (Kural 4.8).'
)

add_table(doc,
    ['Dongu', 'Kumulatif Ay', 'Onceki Toplam', 'Bu Donem Artis', 'Tutar'],
    [
        ['10 Mayis', '1 ay', '0', '1 ay', '1500 TL'],
        ['10 Haziran', '2 ay', '1500 TL (1 ay)', '1 ay', '1500 TL'],
        ['10 Temmuz', '3 ay', '3000 TL (2 ay)', '1 ay', '1500 TL'],
    ]
)

doc.add_paragraph('')
doc.add_paragraph('Her dongude sadece "yeni tamamlanan ay" icin kayit olusturulur. Gecmis kayitlar degismez.')

# ===== 4. HAFTALIK HESAPLAMA DETAYI =====
add_heading(doc, '4. Mod 2: Haftalik Hesaplama (Sonlandirma)', level=1)

add_heading(doc, '4.1 Ne Zaman Kullanilir?', level=2)
doc.add_paragraph('Ogrencinin mentor iliskisi sonlandiginda kullanilir:')

add_table(doc,
    ['Senaryo', 'Tetikleyici', 'Brief Kurali'],
    [
        ['Mentor degisikligi', 'Admin panelde mentor degistir', 'Kural 4.6'],
        ['Ogrenci birakma', 'Ogrenci "birakti" isaretlenir', 'Kural 4.7'],
        ['Iade', 'Ogrenci "iade" isaretlenir', 'Kural 4.7'],
        ['14 gun icinde iade', 'Ogrenci 14 gun icinde iade alir', 'Kural 4.7'],
    ]
)

add_heading(doc, '4.2 Mantik', level=2)
doc.add_paragraph(
    'Sonlandirma aninda, mentorun o ogrenciyle calistigi toplam sureye bakilir. '
    'Sadece TAM 7 gunluk haftalar sayilir. Kismi haftalar = 0 TL.'
)

add_table(doc,
    ['Tamamlanan Sure', 'Hesap', 'Tutar'],
    [
        ['0-6 gun (0 hafta)', '0 x 375', '0 TL'],
        ['7-13 gun (1 hafta)', '1 x 375', '375 TL'],
        ['14-20 gun (2 hafta)', '2 x 375', '750 TL'],
        ['21-27 gun (3 hafta)', '3 x 375', '1125 TL'],
        ['28-34 gun (4 hafta)', '4 x 375', '1500 TL'],
        ['35-41 gun (5 hafta)', '5 x 375', '1875 TL'],
    ]
)

add_heading(doc, '4.3 Adim Adim Ornek: Mentor Degisikligi', level=2)

add_bold_paragraph(doc, 'Ogrenci: Ahmet, SAG: 2 Subat, Mentor Degisikligi: 11 Subat')

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run(
    '2 Subat ████████ 9 Subat  ░░ 11 Subat\n'
    '         |  7 gun  |=1 hafta| 2 gun |\n'
    '         Tamamlandi         Tamamlanmadi'
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph('')
doc.add_paragraph('1. Mentor (Efe):', style='List Bullet')
doc.add_paragraph('Atama: 2 Subat → 11 Subat = 9 gun', style='List Bullet 2')
doc.add_paragraph('Tamamlanan hafta: 1 (9 gun / 7 = 1.28 → 1)', style='List Bullet 2')
doc.add_paragraph('Kazanc: 1 x 375 = 375 TL', style='List Bullet 2')

doc.add_paragraph('')
doc.add_paragraph('2. Yeni Mentor (Ayse):', style='List Bullet')
doc.add_paragraph('Atama: 11 Subat → 2 Mart = 19 gun', style='List Bullet 2')
doc.add_paragraph('Tamamlanan hafta: 2 (19 gun / 7 = 2.71 → 2)', style='List Bullet 2')
doc.add_paragraph('Kazanc: 2 x 375 = 750 TL', style='List Bullet 2')

doc.add_paragraph('')
add_bold_paragraph(doc, 'Toplam: 375 + 750 = 1125 TL (aylik 1500 TL\'nin altinda, cunku ay tamamlanmadi)')

add_heading(doc, '4.4 Adim Adim Ornek: Iade', level=2)

add_bold_paragraph(doc, 'Ogrenci: Fatma, SAG: 3 Mart, Iade: 11 Mart (8 gun sonra)')

doc.add_paragraph('')
doc.add_paragraph('Atama: 3 Mart → 11 Mart = 8 gun', style='List Bullet')
doc.add_paragraph('Tamamlanan hafta: 1 (8 gun / 7 = 1.14 → 1)', style='List Bullet')
doc.add_paragraph('Kazanc: 1 x 375 = 375 TL', style='List Bullet')
doc.add_paragraph('')
doc.add_paragraph('Not: 14 gun icinde %100 iade olsa bile mentor tamamlanan hafta kadar odeme alir (Kural 4.7).')

add_heading(doc, '4.5 Adim Adim Ornek: Kismi Hafta = 0 TL', level=2)

add_bold_paragraph(doc, 'Ogrenci: Zeynep, SAG: 2 Subat, Mentor Degisikligi: 5 Subat (3 gun)')

doc.add_paragraph('')
doc.add_paragraph('Atama: 2 Subat → 5 Subat = 3 gun', style='List Bullet')
doc.add_paragraph('Tamamlanan hafta: 0 (3 gun / 7 = 0.42 → 0)', style='List Bullet')
add_bold_paragraph(doc, 'Kazanc: 0 TL', ' (1 hafta tamamlanmadi)')

# ===== 5. IKI MOD BIR ARADA =====
add_heading(doc, '5. Iki Mod Bir Arada: Senaryo Akislari', level=1)

add_heading(doc, '5.1 Normal Devam Senaryosu', level=2)
add_bold_paragraph(doc, 'Ogrenci kayit olur → 3 ay boyunca devam eder → paket biter')

doc.add_paragraph('')
add_table(doc,
    ['Adim', 'Tarih', 'Olay', 'Hesaplama Modu', 'Tutar'],
    [
        ['1', '10 Nis', 'Ogrenci kayit, mentor atandi', '-', '-'],
        ['2', '10 May', '1. dongu tarihi (1 ay tamamlandi)', 'AYLIK', '1500 TL'],
        ['3', '10 Haz', '2. dongu tarihi (2 ay tamamlandi)', 'AYLIK', '1500 TL'],
        ['4', '10 Tem', '3. dongu tarihi (3 ay tamamlandi)', 'AYLIK', '1500 TL'],
        ['', '', 'TOPLAM', '', '4500 TL'],
    ]
)

add_heading(doc, '5.2 Mentor Degisikligi Senaryosu', level=2)
add_bold_paragraph(doc, 'Ogrenci 2 ay devam eder → 3. ayin ortasında mentor degisir')

doc.add_paragraph('')
add_table(doc,
    ['Adim', 'Tarih', 'Olay', 'Hesaplama Modu', 'Tutar'],
    [
        ['1', '10 Nis', 'Ogrenci kayit, Efe atandi', '-', '-'],
        ['2', '10 May', '1. dongu (Efe)', 'AYLIK', '1500 TL'],
        ['3', '10 Haz', '2. dongu (Efe)', 'AYLIK', '1500 TL'],
        ['4', '25 Haz', 'Mentor degisikligi: Efe → Ayse', '-', '-'],
        ['5', '25 Haz', 'Efe sonlandirma (10 Haz→25 Haz=15 gun=2 hafta)', 'HAFTALIK', '750 TL'],
        ['6', '10 Tem', '3. dongu (Ayse: 25 Haz→10 Tem=15 gun=2 hafta)', 'HAFTALIK', '750 TL'],
        ['', '', 'Efe TOPLAM', '', '1500+1500+750 = 3750 TL'],
        ['', '', 'Ayse TOPLAM', '', '750 TL'],
    ]
)

doc.add_paragraph('')
doc.add_paragraph('Not: Efe\'nin periyodik kazanci (1500+1500) aylik modda hesaplandi. Sonlandirma (750 TL) haftalik modda.', style='List Bullet')
doc.add_paragraph('Ayse\'nin ilk donemi tam ay olmadigi icin haftalik modda hesaplandi.', style='List Bullet')

add_heading(doc, '5.3 Iade Senaryosu', level=2)
add_bold_paragraph(doc, 'Ogrenci 1 ay devam eder → 2. ayin basinda iade alir')

doc.add_paragraph('')
add_table(doc,
    ['Adim', 'Tarih', 'Olay', 'Hesaplama Modu', 'Tutar'],
    [
        ['1', '10 Nis', 'Ogrenci kayit', '-', '-'],
        ['2', '10 May', '1. dongu', 'AYLIK', '1500 TL'],
        ['3', '18 May', 'Iade talebi (8 gun sonra)', '-', '-'],
        ['4', '18 May', 'Sonlandirma (10 May→18 May=8 gun=1 hafta)', 'HAFTALIK', '375 TL'],
        ['', '', 'Mentor TOPLAM', '', '1500 + 375 = 1875 TL'],
    ]
)

# ===== 6. PAKET TURLERI =====
add_heading(doc, '6. Paket Turleri ve Otomatik Bitis Tarihi', level=1)

add_heading(doc, '6.1 Paket Turleri', level=2)
doc.add_paragraph('Ogrenci kayit olurken paket turu secilir. Bu bitis tarihini belirler.')

add_table(doc,
    ['Paket Turu', 'Sure', 'Otomatik Bitis Tarihi', 'Mentor Toplam Kazanc'],
    [
        ['1 Aylik', '1 ay', 'SAG + 1 ay', '1 x 1500 = 1500 TL'],
        ['3 Aylik', '3 ay', 'SAG + 3 ay', '3 x 1500 = 4500 TL'],
        ['6 Aylik', '6 ay', 'SAG + 6 ay', '6 x 1500 = 9000 TL'],
        ['YKS\'ye Kadar', 'Belirsiz', 'Manuel giris zorunlu', 'Her ay 1500 TL (devam ettikce)'],
    ]
)

add_heading(doc, '6.2 Otomatik Bitis Tarihi Hesaplama', level=2)
doc.add_paragraph('Paket turune gore bitis tarihi otomatik hesaplanir:')

doc.add_paragraph('')
add_table(doc,
    ['SAG', 'Paket', 'Bitis Tarihi Hesabi', 'Sonuc'],
    [
        ['10 Nisan', '1 Aylik', '10 Nisan + 1 ay', '10 Mayis'],
        ['10 Nisan', '3 Aylik', '10 Nisan + 3 ay', '10 Temmuz'],
        ['10 Nisan', '6 Aylik', '10 Nisan + 6 ay', '10 Ekim'],
        ['31 Ocak', '1 Aylik', '31 Ocak + 1 ay (Subatta 28 gun)', '28 Subat'],
        ['31 Ocak', '3 Aylik', '31 Ocak + 3 ay', '30 Nisan'],
    ]
)

doc.add_paragraph('')
doc.add_paragraph('Ay sonu tasmasi guvenli: 31 Ocak + 1 ay = 28 Subat ( artik yil = 29 Subat)', style='List Bullet')

# ===== 7. DONGU AKIS SEMASI =====
add_heading(doc, '7. Tam Dongu Akis Semasi', level=1)

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run(
    '═══════════════════════════════════════════════════════════════\n'
    '                    OGRENCI KAYIT OLUR\n'
    '                    SAG = 10 Nisan\n'
    '                    Paket = 3 Aylik\n'
    '                    Bitis = 10 Temmuz (otomatik)\n'
    '═══════════════════════════════════════════════════════════════\n'
    '                              |\n'
    '                    MENTOR ATANIR\n'
    '                              |\n'
    '              ┌───────────────┴───────────────┐\n'
    '              |                               |\n'
    '         DURUM A:                      DURUM B:\n'
    '     Ogrenci devam ediyor          Ogrenci birakma/iade\n'
    '              |                               |\n'
    '     PERIYODIK HESAPLAMA         SONLANDIRMA\n'
    '     (Aylik mod)                 (Haftalik mod)\n'
    '              |                               |\n'
    '     Her dongu tarihi:           Bir kerelik:\n'
    '     + 1500 TL                   Toplam tam hafta\n'
    '              |                   x 375 TL\n'
    '         10 Mayis → 1500 TL               |\n'
    '         10 Haziran → 1500 TL         Final earning\n'
    '         10 Temmuz → 1500 TL           olusturulur\n'
    '              |                               |\n'
    '         TOPLAM: 4500 TL          TOPLAM: Degisen\n'
    '═══════════════════════════════════════════════════════════════'
)
run.font.name = 'Consolas'
run.font.size = Pt(8)

# ===== 8. KOD DEGISIKLIK PLANI =====
add_heading(doc, '8. Kod Degisiklik Plani', level=1)

add_heading(doc, '8.1 Mevcut Durum', level=2)
doc.add_paragraph(
    'Su anda dongu tarihleri SAG bazli aylik olarak uretiliyor (dogru). '
    'Ancak periyodik hesaplamada tutar hala haftalik olarak hesaplaniyor. '
    'Bu, tam aylarda ayni sonucu veriyor (4 hafta x 375 = 1500 TL) ama '
    '31 gunluk aylarda 5 tam hafta cikabilir → 5 x 375 = 1875 TL (hatali).'
)

add_heading(doc, '8.2 Yapilacak Degisiklikler', level=2)

add_table(doc,
    ['Degisiklik', 'Dosya', 'Aciklama'],
    [
        ['Yeni fonksiyon: calculateMonthlyEarning()', 'mentor-earnings.ts', 'Periyodik hesaplama icin: her tamamlanan ay = 1500 TL'],
        ['calculatePendingEarnings guncelleme', 'mentor-earnings.ts', 'Haftalik yerine aylik hesaplama kullanacak'],
        ['calculatePendingEarningsForMentor guncelleme', 'mentor-earnings.ts', 'Ayni degisiklik'],
        ['finalizeMentorEarningForAssignment KORUNUR', 'mentor-earnings.ts', 'Sonlandirma hala haftalik (375 TL/hafta)'],
        ['Paket turleri genisletme', 'schema + API + form', '1/3/6 aylik + YKS kadar secenekleri'],
        ['Otomatik endDate hesaplama', 'API route', 'Paket turune gore bitis tarihi hesaplama'],
    ]
)

add_heading(doc, '8.3 Degismeyecek Kisimlar', level=2)
doc.add_paragraph('Sonlandirma hesabi (finalizeMentorEarningForAssignment): Haftalik hesaplama korunacak', style='List Bullet')
doc.add_paragraph('Dongu tarihi uretimi (getAllCycleDates): SAG bazli aylik, zaten dogru', style='List Bullet')
doc.add_paragraph('Odeme tarihi hesaplama (getPaymentDateForCycle): 1/15 mantigi korunacak', style='List Bullet')
doc.add_paragraph('3 odeme limiti (mentor API): paymentDate bazli sinirlama korunacak', style='List Bullet')
doc.add_paragraph('Rate-lock mekanizmasi: Her kayda weeklyRate/monthlyRate yazilmaya devam edecek', style='List Bullet')

# ===== 9. VERI MODELI =====
add_heading(doc, '9. Veri Modeli Degisiklikleri', level=1)

add_heading(doc, '9.1 MentorEarning Kaydi', level=2)
doc.add_paragraph(
    'Her earning kaydinda hangi hesaplama modunun kullanildigini belirtmek icin '
    'triggerReason alani kullanilir:'
)

add_table(doc,
    ['triggerReason', 'Hesaplama Modu', 'Aciklama'],
    [
        ['periodic_calc', 'AYLIK', 'Periyodik hesaplama (her ay otomatik)'],
        ['assignment_end', 'HAFTALIK', 'Mentor degisikliginde sonlandirma'],
        ['student_drop', 'HAFTALIK', 'Ogrenci birakma'],
        ['student_refund', 'HAFTALIK', 'Iade'],
        ['student_refund_14day', 'HAFTALIK', '14 gun icinde iade'],
        ['membership_renewal', 'HAFTALIK', 'Uyelik yenileme (eski donem sonlandirma)'],
    ]
)

add_heading(doc, '9.2 MentorEarning Alani Degisikligi', level=2)
doc.add_paragraph('Earning kaydinda takip edilen degerler:')

add_table(doc,
    ['Alan', 'Aylik Modda', 'Haftalik Modda'],
    [
        ['completedWeeks', 'Kullanilmiyacak (veya 0)', 'Tamamlanan hafta sayisi'],
        ['completedMonths', '1 (her kayit 1 ay)', 'Kullanilmiyacak'],
        ['amount', '1500 TL (sabit)', 'completedWeeks x 375 TL'],
        ['weeklyRate', '375 TL (referans)', '375 TL (hesaplama icin)'],
        ['monthlyRate', '1500 TL (hesaplama icin)', '1500 TL (referans)'],
    ]
)

doc.add_paragraph('')
doc.add_paragraph('Not: completedMonths alani mevcut modelde yok, eklenecek. Ya da completedWeeks alani her iki modda da kullanilabilir (aylik modda 4 olarak yazilabilir).')

# ===== 10. TEST SENARYOLARI =====
add_heading(doc, '10. Test Senaryolari', level=1)

add_table(doc,
    ['#', 'Senaryo', 'Beklenen Sonuc'],
    [
        ['T1', '1 aylik paket, tam 1 ay devam', '1 x 1500 = 1500 TL (aylik)'],
        ['T2', '3 aylik paket, tam 3 ay devam', '3 x 1500 = 4500 TL (aylik)'],
        ['T3', '1 aylik paket, 10 gun sonra mentor degisikligi', '1 x 375 = 375 TL (haftalik, sonlandirma)'],
        ['T4', '6 aylik paket, 2 ay sonra birakma', '2 x 1500 (aylik) + kismi ay x 375 (haftalik)'],
        ['T5', '14 gun icinde iade', '1 veya 2 hafta x 375 TL (haftalik)'],
        ['T6', 'SAG 31 Ocak, 1 aylik paket', 'Bitis 28 Subat, 1 x 1500 = 1500 TL'],
        ['T7', 'SAG 19 Nisan, 3 aylik paket', 'Donguler: 19 May, 19 Haz, 19 Tem | Odeme: 1 Haz, 1 Tem, 1 Agu'],
        ['T8', 'Mentor panelde 3\'den fazla ogrenci', 'Sadece 3 odeme tarihi gormeli'],
        ['T9', 'Ogrenci yenileme yapar', 'Eski donem sonlandirma (haftalik) + yeni donem aylik'],
        ['T10', 'YKS\'ye kadar paket', 'Suresiz dongu, her ay 1500 TL'],
    ]
)

# Kaydet
path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Coachify_Odeme_Mantigi_Detay.docx')
doc.save(path)
print('Dosya olusturuldu: ' + path)
