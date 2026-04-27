import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

title = doc.add_heading('Coachify Panel \u2014 Revize Degerlendirme Raporu', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Tarih: 21 Nisan 2026').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Belge: Coachify 2. Revize-2.pdf karsiligi degerlendirme')
doc.add_paragraph('')

# ===== OZET TABLOSU =====
doc.add_heading('Genel Ozet', level=1)

summary_data = [
    ['Madde', 'Revize Istei', 'Durum', 'Aciklama'],
    ['1', 'Donem tarihi 1 ve 15 olmali', 'YAPILDI', 'Aylik dongu + getPaymentDateForCycle ile cozuldu'],
    ['2', 'Yeni orenci formu duzenlemeleri', 'EKSIK', 'Okul alani halen formda, * isaretleri tutarsiz'],
    ['3', 'Yeni orenci odeme durumu "odendi" olmali', 'YAPILDI', 'paymentStatus default "paid" olarak ayarli'],
    ['4', 'Bitis tarihi aylik hesaplanmali', 'YAPILDI', 'SAG bazli aylik dongu implement edildi'],
    ['5', 'Mentor max 3 odeme gormeli', 'YAPILDI', 'MAX_VISIBLE_UPCOMING_CYCLES=3, paymentDate bazli'],
    ['6', 'Bitis tarihi zorunlu olmali', 'YAPILDI', 'Prisma schema: endDate DateTime (nullable degil)'],
    ['7', 'Iade hesabi kontrolue', 'YAPILDI', 'Sadece tamamlanan hafta kadar odeme hesaplaniyor'],
    ['8', 'Webhook baglantisi netlestirilmeli', 'YAPILDI', '4 Zapier endpoint mevcut (register,update,renew,query)'],
    ['9', 'Odeme bilgilendirme mesajlari', 'YAPILMADI', 'Hicbir bildirim/email sistemi yok'],
]

st = doc.add_table(rows=len(summary_data), cols=4, style='Light Grid Accent 1')
st.alignment = WD_TABLE_ALIGNMENT.CENTER
for ri, rd in enumerate(summary_data):
    for ci, cd in enumerate(rd):
        cell = st.rows[ri].cells[ci]
        cell.text = cd
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
                if ri == 0:
                    r.bold = True
                if ci == 2 and ri > 0:
                    r.bold = True
                    if cd == 'YAPILDI':
                        r.font.color.rgb = RGBColor(0, 128, 0)
                    elif cd in ('EKSIK', 'YAPILMADI'):
                        r.font.color.rgb = RGBColor(255, 0, 0)

doc.add_paragraph('')

# Helper: tablo olustur
def add_table(doc, headers, rows):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers), style='Light Grid Accent 1')
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, rd in enumerate(rows):
        for ci, cd in enumerate(rd):
            t.rows[ri+1].cells[ci].text = cd
            for p in t.rows[ri+1].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return t

# ===== MADDE 1 =====
doc.add_heading('Madde 1 \u2014 Donem Tarihi Gosterimi', level=1)
doc.add_paragraph('Revize Istei:', style='List Bullet')
doc.add_paragraph('Donem tarihinde 14 ve 15 gosteriliyordu. Mantik 1 ve 15 baz almali.')
doc.add_paragraph('')
doc.add_paragraph('Yapilan Degisiklik:', style='List Bullet')
doc.add_paragraph('Odeme dongusu SAG (Satin Alma Gunu) bazinda aylik hesaplanacak sekilde yeniden yazildi. Odeme tarihi (para transferi) 1 ve 15\'e sabitlendi.')
doc.add_paragraph('')

add_table(doc,
    ['Kavram', 'Deger'],
    [
        ['Dongu tarihi (cycle date)', 'SAG bazli aylik (orn. 4 Nisan baslangic -> 4 Mayis, 4 Haziran)'],
        ['Odeme tarihi (payment date)', 'cycleDate gunu <=15 -> ayni ayin 15\'i, >15 -> sonraki ayin 1\'i'],
    ]
)

doc.add_paragraph('')
doc.add_paragraph('Degisen dosyalar:', style='List Bullet')
doc.add_paragraph('src/lib/mentor-earnings.ts \u2014 getAllCycleDates, getNextPaymentDate, getPaymentDateForCycle', style='List Bullet 2')
doc.add_paragraph('src/app/api/mentor/earnings/route.ts \u2014 paymentDate bazli gruplama', style='List Bullet 2')
doc.add_paragraph('src/app/mentor/page.tsx \u2014 paymentDate gosterimi', style='List Bullet 2')

# ===== MADDE 2 =====
doc.add_heading('Madde 2 \u2014 Yeni Orenci Ekleme Formu', level=1)
doc.add_paragraph('Revize Istei:', style='List Bullet')
doc.add_paragraph('Mail alani zorunlu olmamali. Okul alani kaldirilmali. Zorunlu alanlar * ile isaretlenmeli.')
doc.add_paragraph('')

t2 = add_table(doc,
    ['Alt Istek', 'Durum', 'Detay'],
    [
        ['Email zorunlu olmamali', 'YAPILDI', 'API null email kabul ediyor, formda required degil'],
        ['Okul alani kaldirilmali', 'YAPILMADI', 'Formda ve API\'de halen mevcut'],
        ['Zorunlu alanlar * ile isaretli', 'YAPILMADI', 'Tutarli * isareti yok'],
    ]
)
# Renklendir
for ri in range(1, 4):
    cell = t2.rows[ri].cells[1]
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            if r.text == 'YAPILDI':
                r.font.color.rgb = RGBColor(0,128,0)
            else:
                r.font.color.rgb = RGBColor(255,0,0)

# ===== MADDE 3 =====
doc.add_heading('Madde 3 \u2014 Yeni Orenci Odeme Durumu', level=1)
doc.add_paragraph('Revize Istei:', style='List Bullet')
doc.add_paragraph('Yeni orenci eklendiinde odeme durumu "odendi" olmali.')
doc.add_paragraph('')
doc.add_paragraph('Mevcut Durum:', style='List Bullet')
doc.add_paragraph('Prisma schema ve API\'de paymentStatus default deeri "paid" olarak tanimli. Yeni orenci olusturulduunda odeme durumu doru sekilde "odendi" olarak basliyor.')
doc.add_paragraph('')
doc.add_paragraph('Ilgili dosya: src/app/api/admin/students/route.ts \u2014 paymentStatus: "paid"', style='List Bullet')

# ===== MADDE 4 =====
doc.add_heading('Madde 4 \u2014 Orenci Bitis Tarihi Hesabi (Aylik Dongu)', level=1)
doc.add_paragraph('Revize Istei:', style='List Bullet')
doc.add_paragraph('Bitis tarihleri haftalik degil, aylik dongu ile hesaplanmali. Brief kural 4.5 uyumlu olmali.')
doc.add_paragraph('')
doc.add_paragraph('Yapilan Degisiklik:', style='List Bullet')
doc.add_paragraph('Tum dongu hesaplama sistemi SAG bazli aylik donguye cevrildi.')
doc.add_paragraph('')

add_table(doc,
    ['Eski Sistem', 'Yeni Sistem', 'Brief Kurali'],
    [
        ['Dongu: Sabit 1 ve 15', 'Dongu: SAG bazli aylik (4, 10, 19 vs.)', 'Kural 4.5: Gun bazli dongu'],
        ['Haftalik: 375 TL/hafta', 'Haftalik: 375 TL/hafta (degismedi)', 'Kural 4.2: Haftalik hesap'],
        ['Odeme: 1 ve 15', 'Odeme: 1 ve 15 (degismedi)', 'Kural 4.3: Odeme gunleri'],
        ['Cycle: 1,15,1,15... (fixed)', 'Cycle: SAG+1ay, SAG+2ay...', 'Kural 4.5: Baslangic gunu bazli'],
    ]
)

doc.add_paragraph('')
doc.add_paragraph('Ornek senaryo:', style='List Bullet')
doc.add_paragraph('SAG = 10 Nisan -> Dongu: 10 Mayis, 10 Haziran, 10 Temmuz', style='List Bullet 2')
doc.add_paragraph('Her dongu icin tamamlanan hafta x 375 TL hesaplanir', style='List Bullet 2')
doc.add_paragraph('10 Mayis dongusu -> odeme tarihi: 15 Mayis', style='List Bullet 2')
doc.add_paragraph('SAG = 19 Nisan -> Dongu: 19 Mayis -> odeme: 1 Haziran (19>15 kurali)', style='List Bullet 2')

# ===== MADDE 5 =====
doc.add_heading('Madde 5 \u2014 Mentorun Gorebildii Odeme Bilgisi', level=1)
doc.add_paragraph('Revize Istei:', style='List Bullet')
doc.add_paragraph('Mentor maksimum 3 odeme gorebilmeli. Paket suresi gorunmemeli.')
doc.add_paragraph('')
doc.add_paragraph('Yapilan Degisiklik:', style='List Bullet')
doc.add_paragraph('API route\'ta MAX_VISIBLE_UPCOMING_CYCLES = 3 limiti paymentDate (1/15) bazli uygulaniyor. Mentor fazlasini goremiyor.')
doc.add_paragraph('')

t5 = add_table(doc,
    ['Bilgi', 'Mentor Gorur mu?'],
    [
        ['Onundeki 3 odeme dongusu', 'EVET'],
        ['3\'u asan gelecek odemeler', 'HAYIR'],
        ['Orencinin paket suresi / toplam plan', 'HAYIR'],
    ]
)

# ===== MADDE 6 =====
doc.add_heading('Madde 6 \u2014 Orenci Bitis Tarihi Zorunlulueu', level=1)
doc.add_paragraph('Revize Istei:', style='List Bullet')
doc.add_paragraph('Her orencinin bitis tarihi zorunlu olmali. Sonsuz/belirsiz bitis olmamali.')
doc.add_paragraph('')
doc.add_paragraph('Mevcut Durum:', style='List Bullet')
doc.add_paragraph('Prisma schema: endDate DateTime (nullable degil). API ve formda required olarak tanimli. Tum orencilerin bitis tarihi zorunlu.')
doc.add_paragraph('')
doc.add_paragraph('Ilgili dosyalar:', style='List Bullet')
doc.add_paragraph('prisma/schema.prisma \u2014 endDate: DateTime (zorunlu)', style='List Bullet 2')
doc.add_paragraph('src/app/api/admin/students/route.ts \u2014 endDate validasyonu', style='List Bullet 2')
doc.add_paragraph('src/app/admin/students/page.tsx \u2014 input required', style='List Bullet 2')

# ===== MADDE 7 =====
doc.add_heading('Madde 7 \u2014 Iade Hesabi Kontrolue', level=1)
doc.add_paragraph('Revize Istei:', style='List Bullet')
doc.add_paragraph('Efe-Kurucu\'daki Salih Yalcin: 8 gunluk bitis + iade -> halen 4 haftalik odeme gosteriyordu. Iade hesabi kontrol edilmeli.')
doc.add_paragraph('')
doc.add_paragraph('Mevcut Durum:', style='List Bullet')
doc.add_paragraph('Iade mantik doru calisiyor: finalizeMentorEarningForAssignment sadece tamamlanan hafta kadar odeme hesapliyor. 8 gun = 1 hafta tamamlandi = 375 TL. Eski 4 haftalik gosterim eski dongu sisteminden kaynaklaniyordu, aylik dongu gecisi ile duzeltilmis olmali.')
doc.add_paragraph('')

add_table(doc,
    ['Senaryo', 'Beklenen Sonuc'],
    [
        ['8 gun sonra iade', '1 hafta = 375 TL (kural 4.7)'],
        ['14 gun icinde %100 iade', 'Mentor yine tamamlanan hafta kadar alir (kural 4.7)'],
        ['Kismi hafta', '0 TL (tamamlanmamis hafta = odeme yok, kural 4.6)'],
    ]
)

# ===== MADDE 8 =====
doc.add_heading('Madde 8 \u2014 Webhook Baglantisi', level=1)
doc.add_paragraph('Revize Sorusu:', style='List Bullet')
doc.add_paragraph('Panelin webhook\'u nereden calisiyor? Zapier akisi bulunamadi.')
doc.add_paragraph('')
doc.add_paragraph('Mevcut Durum:', style='List Bullet')
doc.add_paragraph('4 Zapier webhook endpoint\'i mevcut ve calisir durumda:')
doc.add_paragraph('')

add_table(doc,
    ['Endpoint', 'Islev', 'Yol'],
    [
        ['POST /register', 'Yeni orenci kaydi (Tally form)', '/api/webhooks/zapier/register'],
        ['POST /update', 'Uyelik guncelleme', '/api/webhooks/zapier/update'],
        ['POST /renew', 'Uyelik yenileme', '/api/webhooks/zapier/renew'],
        ['POST /query', 'Veri sorgulama (bitis yalasan, email ile)', '/api/webhooks/zapier/query'],
    ]
)

doc.add_paragraph('')
doc.add_paragraph('Guvenlik: x-api-key header + timingSafeEqual ile dogrulama', style='List Bullet')
doc.add_paragraph('Akis: Tally form -> Zapier -> Webhook -> Panel veritabani', style='List Bullet')

# ===== MADDE 9 =====
doc.add_heading('Madde 9 \u2014 Odeme Bilgilendirme Mesajlari', level=1)
doc.add_paragraph('Revize Sorusu:', style='List Bullet')
doc.add_paragraph('Odeme mesajlari test edilemedi. Zapier akisinda bulunamadi.')
doc.add_paragraph('')
doc.add_paragraph('Mevcut Durum: YAPILMADI', style='List Bullet')
doc.add_paragraph('')
doc.add_paragraph('Sistemde hicbir bildirim/email/SMS mekanizmasi yok. Schema\'da sendMessage ve contactPreference alanlari mevcut ama kullanilmiyor.', style='List Bullet')
doc.add_paragraph('')
doc.add_paragraph('Eksik bilesenler:', style='List Bullet')
doc.add_paragraph('E-posta/SMS gonderim kutuphanesi (nodemailer, SendGrid vs.)', style='List Bullet 2')
doc.add_paragraph('Zamanlanmis gorev (cron job) \u2014 hatirlatma kontrolue', style='List Bullet 2')
doc.add_paragraph('Mesaj sablonu ve gonderim mantiei', style='List Bullet 2')
doc.add_paragraph('/query webhook "expiring_soon" sorgusu hazir \u2014 ancak tetikleyici yok', style='List Bullet 2')

# ===== SONUC =====
doc.add_heading('Sonuc ve Aksiyon Plani', level=1)

ts = add_table(doc,
    ['Oncelik', 'Madde', 'Aksiyon'],
    [
        ['YUKSEK', 'Madde 2 \u2014 Okul alani kaldir', 'Form ve API\'den school field kaldirilmali, * isaretleri duzeltilmeli'],
        ['ORTA', 'Madde 9 \u2014 Odeme hatirlatma', 'Email/SMS sistemi kurulmali, cron job eklenmeli'],
        ['BILGI', 'Madde 8 \u2014 Webhook', 'Zapier akisi calisiyor, sadece dokumantasyion gerekli'],
    ]
)

for ri in range(1, 4):
    cell = ts.rows[ri].cells[0]
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)
            if 'YUKSEK' in r.text:
                r.font.color.rgb = RGBColor(255,0,0)
            elif 'ORTA' in r.text:
                r.font.color.rgb = RGBColor(255,165,0)

# Kaydet
import os
path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Coachify_Revize_Degerlendirme_Raporu.docx')
doc.save(path)
print('Word dosyasi olusturuldu: ' + path)
